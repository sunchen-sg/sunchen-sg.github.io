from __future__ import annotations

import argparse
import re
import unicodedata
from pathlib import Path

from docx import Document
from docx.enum.text import WD_UNDERLINE


TITLE_RE = re.compile(r"“(.+?)”")
ROLE_RE = re.compile(r"\.\s+\(([^()]+)\)\s*$")
MARKER_RE = re.compile(r"^(.*?)([*#†]+)?$")

def normalized_title(value: str) -> str:
    value = unicodedata.normalize("NFKD", value)
    value = value.replace("–", "-").replace("—", "-")
    value = re.sub(r"[^a-z0-9]+", " ", value.lower())
    return " ".join(value.split())


# The current homepage retains the verified co-first-author annotation for this
# paper, while the latest SelectedPublications.docx export omits it.
VERIFIED_SELECTED_OVERRIDES = {
    normalized_title(
        "Enabling UTBB strained SOI platform for co-integration of logic and RF: "
        "Implant-induced strain relaxation and comb-like device architecture"
    ): {
        "markers": {"C. Sun": "*", "J. Liang": "*"},
        "role": "co-first author",
    }
}


def citation_title(text: str) -> str | None:
    match = TITLE_RE.search(text)
    return match.group(1) if match else None


def author_markers(author_text: str) -> dict[str, str]:
    author_text = author_text.replace(", and ", ", ").replace(" and ", ", ")
    markers: dict[str, str] = {}
    for token in (part.strip() for part in author_text.split(",")):
        if not token:
            continue
        match = MARKER_RE.match(token)
        if not match:
            continue
        name = match.group(1).strip()
        marker = match.group(2) or ""
        if marker:
            markers[name] = marker
    return markers


def selected_metadata(path: Path) -> dict[str, dict[str, object]]:
    document = Document(path)
    selected: dict[str, dict[str, object]] = {}
    for paragraph in document.paragraphs:
        if paragraph.text.lstrip().startswith("(Granted US Patent)"):
            continue
        title = citation_title(paragraph.text)
        if not title:
            continue
        author_text = paragraph.text.split(", “", 1)[0]
        role_match = ROLE_RE.search(paragraph.text)
        selected[normalized_title(title)] = {
            "title": title,
            "markers": author_markers(author_text),
            "role": role_match.group(1) if role_match else None,
            "to_appear": "to appear in" in paragraph.text.lower(),
        }
    for key, override in VERIFIED_SELECTED_OVERRIDES.items():
        if key not in selected:
            raise ValueError(f"Verified selected override does not match a selected paper: {key}")
        selected[key].update(override)
    return selected


def append_marker(paragraph, author_name: str, marker: str, quote_index: int) -> bool:
    for run in paragraph.runs[:quote_index]:
        if author_name not in run.text:
            continue
        if f"{author_name}{marker}" in run.text:
            return True
        run.text = run.text.replace(author_name, f"{author_name}{marker}", 1)
        return True
    return False


def move_title_comma_inside(paragraph, quote_index: int) -> None:
    runs = paragraph.runs
    close_index = next(
        (index for index in range(quote_index, len(runs)) if "”" in runs[index].text),
        None,
    )
    if close_index is None:
        raise ValueError(f"Missing closing quotation mark: {paragraph.text}")

    close_run = runs[close_index]
    if re.search(r"”\s*,", close_run.text):
        close_run.text = re.sub(r"”\s*,", ",”", close_run.text, count=1)
    elif ",”" not in close_run.text:
        close_run.text = close_run.text.replace("”", ",”", 1)

    for run in runs[close_index + 1 :]:
        if not run.text:
            continue
        run.text = re.sub(r"^,\s*", " ", run.text, count=1)
        break


def mark_to_appear(paragraph, quote_index: int) -> None:
    for run in paragraph.runs[quote_index:]:
        if "to appear in" in run.text.lower():
            return
        if re.search(r"\bin\s+$", run.text):
            run.text = re.sub(r"\bin\s+$", "to appear in ", run.text, count=1)
            return
        if run.text.startswith("in "):
            run.text = f"to appear {run.text}"
            return
    raise ValueError(f"Could not add to-appear status: {paragraph.text}")


def emphasize_sun(paragraph, quote_index: int) -> None:
    matches = 0
    for run in paragraph.runs[:quote_index]:
        if "C. Sun" not in run.text:
            continue
        run.bold = True
        run.underline = WD_UNDERLINE.SINGLE
        matches += 1
    if matches != 1:
        raise ValueError(f"Expected one C. Sun author token, found {matches}: {paragraph.text}")


def append_role(paragraph, role: str) -> None:
    if f"({role})" in paragraph.text:
        return
    paragraph.add_run(" (")
    role_run = paragraph.add_run(role)
    role_run.bold = True
    paragraph.add_run(")")


def normalize_document(source: Path, selected_source: Path, output: Path) -> None:
    selected = selected_metadata(selected_source)
    if len(selected) != 20:
        raise ValueError(f"Expected 20 selected publications, found {len(selected)}")

    document = Document(source)
    citation_count = 0
    selected_matches: set[str] = set()
    marker_updates = 0
    role_updates = 0

    for paragraph in document.paragraphs:
        title = citation_title(paragraph.text)
        if not title:
            continue
        citation_count += 1
        key = normalized_title(title)
        quote_index = next(
            (index for index, run in enumerate(paragraph.runs) if "“" in run.text),
            None,
        )
        if quote_index is None:
            raise ValueError(f"Missing opening quotation mark: {paragraph.text}")

        move_title_comma_inside(paragraph, quote_index)
        emphasize_sun(paragraph, quote_index)

        metadata = selected.get(key)
        if not metadata:
            continue
        selected_matches.add(key)
        for author_name, marker in metadata["markers"].items():
            if not append_marker(paragraph, author_name, marker, quote_index):
                raise ValueError(
                    f"Could not attach {marker!r} to {author_name!r}: {paragraph.text}"
                )
            marker_updates += 1

        if metadata["to_appear"]:
            mark_to_appear(paragraph, quote_index)
        role = metadata["role"]
        if role:
            append_role(paragraph, str(role))
            role_updates += 1

    missing = set(selected) - selected_matches
    if missing:
        missing_titles = [str(selected[key]["title"]) for key in sorted(missing)]
        raise ValueError(f"Selected publications missing from All Publications: {missing_titles}")
    if citation_count != 103:
        raise ValueError(f"Expected 103 citation/patent entries, found {citation_count}")

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    print(f"Normalized citations: {citation_count}")
    print(f"Selected publications matched: {len(selected_matches)}")
    print(f"Contribution marker attachments checked: {marker_updates}")
    print(f"Role annotations checked: {role_updates}")
    print(f"Output: {output}")


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--selected", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    normalize_document(args.source, args.selected, args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
