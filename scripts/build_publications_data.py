from __future__ import annotations

import argparse
import difflib
import html
import json
import re
import unicodedata
import urllib.parse
import urllib.request
from collections import defaultdict
from pathlib import Path

from docx import Document


CO_COLLECTION_KEY = "JGVR7XDM"
APPROVED_TO_APPEAR_KEY = "2MF3SQGT"
VERIFIED_DOI_OVERRIDES = {
    "PGSW855G": "10.1109/IEDM50572.2025.11353857",
    "JXPBVZXS": "10.1109/IEDM50572.2025.11353726",
}
VERIFIED_ITEM_TYPE_OVERRIDES = {
    # Official 242nd ECS Meeting program identifies this ECS Meeting Abstract as a conference paper.
    "LCHXJI6L": "conferencePaper",
}
TITLE_RE = re.compile(r"“(.+?)[,]?”")
ROLE_RE = re.compile(r"\.\s+\(([^()]+)\)\s*$")
SUN_RE = re.compile(r"C\. Sun([*#†]*)")
YEAR_RE = re.compile(r"20\d{2}")


def normalized_title(value: str) -> str:
    value = re.sub(r"<[^>]+>", "", html.unescape(value))
    value = unicodedata.normalize("NFKD", value)
    value = value.replace("–", "-").replace("—", "-")
    value = re.sub(r"[^a-z0-9]+", " ", value.lower())
    return " ".join(value.split())


def citation_title(text: str) -> str | None:
    match = TITLE_RE.search(text)
    return match.group(1) if match else None


def selected_titles(path: Path) -> set[str]:
    document = Document(path)
    titles = {
        normalized_title(title)
        for paragraph in document.paragraphs
        if not paragraph.text.lstrip().startswith("(Granted US Patent)")
        if (title := citation_title(paragraph.text))
    }
    if len(titles) != 20:
        raise ValueError(f"Expected 20 selected publications, found {len(titles)}")
    return titles


def paragraph_venue(paragraph) -> str:
    close_seen = False
    venue_runs: list[str] = []
    for run in paragraph.runs:
        if "”" in run.text:
            close_seen = True
            continue
        if close_seen and run.italic:
            venue_runs.append(run.text)
    return "".join(venue_runs).strip()


def parse_publications(path: Path, selected: set[str]) -> list[dict[str, object]]:
    document = Document(path)
    records: list[dict[str, object]] = []
    year: int | None = None

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text == "Patents":
            break
        if YEAR_RE.fullmatch(text):
            year = int(text)
            continue
        title = citation_title(text)
        if year is None or not title:
            continue

        venue = paragraph_venue(paragraph)
        if not venue:
            raise ValueError(f"Missing italic publication venue: {text}")
        title_marker = f"“{title},”"
        if title_marker not in text:
            raise ValueError(f"Citation title punctuation is not normalized: {text}")

        authors = text.split(", “", 1)[0]
        sun_match = SUN_RE.search(authors)
        if not sun_match:
            raise ValueError(f"Missing C. Sun author token: {text}")

        tail = text.split(title_marker, 1)[1].strip()
        venue_start = tail.find(venue)
        if venue_start < 0:
            raise ValueError(f"Could not locate venue in citation: {text}")
        lead = tail[:venue_start].strip()
        details = tail[venue_start + len(venue) :].strip()
        role_match = ROLE_RE.search(text)
        role = role_match.group(1) if role_match else None
        if role:
            details = re.sub(r"\s+\([^()]+\)\s*$", "", details)

        records.append(
            {
                "year": year,
                "authorsPrefix": authors[: sun_match.start()],
                "sunAuthor": sun_match.group(0),
                "authorsSuffix": authors[sun_match.end() :],
                "title": title,
                "lead": lead,
                "venue": venue,
                "details": details,
                "role": role,
                "selectedBySource": normalized_title(title) in selected,
                "sourceOrder": len(records),
            }
        )

    if len(records) != 102:
        raise ValueError(f"Expected 102 journal/conference publications, found {len(records)}")
    return records


def fetch_zotero_items(api_base: str) -> list[dict[str, object]]:
    items: list[dict[str, object]] = []
    start = 0
    while True:
        query = urllib.parse.urlencode(
            {
                "limit": 100,
                "start": start,
                "include": "data",
                "sort": "title",
                "direction": "asc",
            }
        )
        request = urllib.request.Request(
            f"{api_base.rstrip('/')}/api/users/0/items/top?{query}",
            headers={"Zotero-API-Version": "3"},
        )
        with urllib.request.urlopen(request, timeout=30) as response:
            page = json.load(response)
        items.extend(page)
        if len(page) < 100:
            break
        start += 100
    return items


def item_year(item: dict[str, object]) -> int | None:
    meta = item.get("meta") or {}
    data = item.get("data") or {}
    date = str(meta.get("parsedDate") or data.get("date") or "")
    match = re.search(r"(?:19|20)\d{2}", date)
    return int(match.group(0)) if match else None


def publication_candidates(items: list[dict[str, object]]) -> list[dict[str, object]]:
    candidates: list[dict[str, object]] = []
    for item in items:
        data = item.get("data") or {}
        collections = set(data.get("collections") or [])
        if CO_COLLECTION_KEY not in collections and item.get("key") != APPROVED_TO_APPEAR_KEY:
            continue
        if data.get("itemType") not in {"journalArticle", "conferencePaper"}:
            continue
        if not data.get("title"):
            continue
        candidates.append(item)
    return candidates


def match_item(record: dict[str, object], candidates: list[dict[str, object]]) -> tuple[dict[str, object], float]:
    target = normalized_title(str(record["title"]))
    exact = [
        item
        for item in candidates
        if normalized_title(str((item.get("data") or {}).get("title") or "")) == target
    ]
    if exact:
        desired_type = (
            "conferencePaper"
            if str(record.get("lead") or "") in {"in", "to appear in"}
            else "journalArticle"
        )
        type_matches = [
            item for item in exact if (item.get("data") or {}).get("itemType") == desired_type
        ]
        exact = type_matches or exact
        year_matches = [item for item in exact if item_year(item) == record["year"]]
        pool = year_matches or exact
        if len(pool) == 1:
            return pool[0], 1.0

    ranked: list[tuple[float, dict[str, object]]] = []
    for item in candidates:
        candidate_title = normalized_title(str((item.get("data") or {}).get("title") or ""))
        score = difflib.SequenceMatcher(None, target, candidate_title).ratio()
        if item_year(item) == record["year"]:
            score += 0.03
        ranked.append((score, item))
    ranked.sort(key=lambda pair: pair[0], reverse=True)
    return ranked[0][1], ranked[0][0]


def link_for(key: str, data: dict[str, object]) -> tuple[str | None, str | None, str | None]:
    doi = VERIFIED_DOI_OVERRIDES.get(key, str(data.get("DOI") or "").strip())
    doi = re.sub(r"^https?://(?:dx\.)?doi\.org/", "", doi, flags=re.IGNORECASE)
    if doi:
        return doi, f"https://doi.org/{doi}", "DOI"
    url = str(data.get("url") or "").strip()
    if url.startswith(("https://", "http://")):
        return None, url, "Publisher"
    return None, None, None


def attach_zotero_identity(
    records: list[dict[str, object]], candidates: list[dict[str, object]]
) -> list[dict[str, object]]:
    used_keys: set[str] = set()
    selected_keys: set[str] = set()
    weakest: list[tuple[float, str, str]] = []

    for record in records:
        item, score = match_item(record, candidates)
        data = item.get("data") or {}
        key = str(item.get("key") or "")
        if score < 0.90:
            raise ValueError(
                f"Low-confidence Zotero match ({score:.3f}): {record['title']} => {data.get('title')}"
            )
        if key in used_keys:
            raise ValueError(f"Duplicate Zotero identity matched: {key} ({record['title']})")
        used_keys.add(key)
        doi, url, link_label = link_for(key, data)
        record.update(
            {
                "id": key,
                "doi": doi,
                "url": url,
                "linkLabel": link_label,
                "itemType": VERIFIED_ITEM_TYPE_OVERRIDES.get(key, data.get("itemType")),
                "selected": bool(record.pop("selectedBySource")),
            }
        )
        if record["selected"]:
            selected_keys.add(key)
        weakest.append((score, str(record["title"]), str(data.get("title") or "")))

    if len(selected_keys) != 20:
        raise ValueError(f"Expected 20 selected stable identities, found {len(selected_keys)}")

    weakest.sort(key=lambda row: row[0])
    print(f"Zotero publication candidates in Co: {len(candidates)}")
    print(f"Matched website publications: {len(records)}")
    print(f"Selected stable identities: {len(selected_keys)}")
    print("Lowest title-match scores:")
    for score, source, zotero in weakest[:8]:
        print(f"  {score:.3f}\t{source}\t=>\t{zotero}")
    print("Selected keys:")
    print("  " + ", ".join(sorted(selected_keys)))
    return records


def write_javascript(records: list[dict[str, object]], output: Path) -> None:
    for record in records:
        record.pop("sourceOrder", None)
    payload = json.dumps(records, ensure_ascii=False, indent=2)
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(
        "// Generated from MyInfo/AllPublications.docx and verified Zotero identities.\n"
        f"window.PUBLICATIONS_DATA = Object.freeze({payload});\n",
        encoding="utf-8",
        newline="\n",
    )
    counts: dict[int, int] = defaultdict(int)
    for record in records:
        counts[int(record["year"])] += 1
    print(f"Output: {output}")
    print("Year counts: " + ", ".join(f"{year}={counts[year]}" for year in sorted(counts, reverse=True)))


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", type=Path, required=True)
    parser.add_argument("--selected", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--zotero-api", default="http://127.0.0.1:23119")
    args = parser.parse_args()

    selected = selected_titles(args.selected)
    records = parse_publications(args.docx, selected)
    items = fetch_zotero_items(args.zotero_api)
    candidates = publication_candidates(items)
    attach_zotero_identity(records, candidates)
    write_javascript(records, args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
